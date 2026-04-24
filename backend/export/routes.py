"""
DOCX export endpoint.

Converts Markdown-formatted AI answers into proper Word documents:
- ## headings → Word Heading styles
- bullet lists → Word List Bullet style
- tables (markdown pipes) → real Word table objects
- bold/italic inline → Word runs with bold/italic
- paragraphs → body text

Use case: users download the full AI answer as a .docx
and paste / open in Google Docs, Word, LibreOffice with
proper formatting — not markdown pipes.
"""
import io
import re
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from backend.auth.routes import get_current_user
from backend.db.models import User

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    answer: str
    question: Optional[str] = ""
    sources: Optional[list] = []


def _parse_inline(text: str):
    """Strip markdown bold/italic markers and return plain text for docx runs."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def _is_bold(text: str) -> bool:
    return bool(re.match(r'^\*\*.+\*\*$', text.strip()))


def _add_run(para, text: str):
    """Add a run to paragraph, applying bold/italic based on markdown markers."""
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for part in parts:
        if re.match(r'\*\*[^*]+\*\*', part):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif re.match(r'\*[^*]+\*', part):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'`[^`]+`', part):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
        elif part:
            para.add_run(part)


def answer_to_docx(answer: str, question: str = "", sources: list = None) -> bytes:
    """
    Convert markdown answer string to DOCX bytes.
    Returns raw bytes suitable for HTTP response.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # Title
    if question:
        title = doc.add_heading(question.strip(), level=1)
        title.runs[0].font.color.rgb = RGBColor(0x09, 0x09, 0x0B)

    # Sources line
    if sources:
        src_para = doc.add_paragraph()
        src_run = src_para.add_run(f"Source(s): {', '.join(sources)}")
        src_run.italic = True
        src_run.font.size = Pt(10)
        src_run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)
        doc.add_paragraph()  # spacer

    lines = answer.split("\n")
    i = 0
    table_buf = []
    in_table = False

    def flush_table():
        nonlocal table_buf
        if len(table_buf) < 2:
            table_buf = []
            return
        heads = [h.strip() for h in table_buf[0].split("|") if h.strip()]
        data_rows = [
            [c.strip() for c in row.split("|") if c.strip()]
            for row in table_buf[2:]  # skip separator row
            if not re.match(r'^[\|\s\-:]+$', row)
        ]
        if not heads:
            table_buf = []
            return

        n_cols = len(heads)
        n_rows = len(data_rows)
        if n_rows == 0:
            table_buf = []
            return

        tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
        tbl.style = 'Table Grid'

        # Header row
        hdr_cells = tbl.rows[0].cells
        for j, h in enumerate(heads):
            if j < len(hdr_cells):
                hdr_cells[j].text = _parse_inline(h)
                for run in hdr_cells[j].paragraphs[0].runs:
                    run.bold = True

        # Data rows
        for ri, row in enumerate(data_rows):
            cells = tbl.rows[ri + 1].cells
            for ci, val in enumerate(row):
                if ci < len(cells):
                    cells[ci].text = _parse_inline(val)

        doc.add_paragraph()  # spacer after table
        table_buf = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Table detection
        if line.strip().startswith("|") and "|" in line:
            in_table = True
            table_buf.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Headings ## / ###
        h_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if h_match:
            level = min(len(h_match.group(1)), 4)
            content = _parse_inline(h_match.group(2))
            doc.add_heading(content, level=level)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*]{3,}$', line.strip()):
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Bullet list
        b_match = re.match(r'^(\s*)([-•*])\s+(.+)', line)
        if b_match:
            indent = len(b_match.group(1))
            content = b_match.group(3)
            style = 'List Bullet 2' if indent >= 2 else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style)
            except Exception:
                p = doc.add_paragraph()
                p.add_run('• ')
            _add_run(p, content)
            i += 1
            continue

        # Numbered list
        n_match = re.match(r'^(\s*)(\d+)[.)]\s+(.+)', line)
        if n_match:
            content = n_match.group(3)
            try:
                p = doc.add_paragraph(style='List Number')
            except Exception:
                p = doc.add_paragraph()
                p.add_run(f"{n_match.group(2)}. ")
            _add_run(p, content)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_run(p, line.strip())
        p.paragraph_format.space_after = Pt(4)
        i += 1

    if in_table:
        flush_table()

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.post("/docx")
def export_answer_docx(
    req: ExportRequest,
    current_user: User = Depends(get_current_user)
):
    """
    Convert an AI answer (markdown) to a downloadable DOCX file.
    The DOCX contains:
    - Question as title
    - Headings, bullets, numbered lists preserved
    - Tables as real Word tables (NOT markdown pipes)
    - Source citations in italics

    Paste in Google Docs: File → Open → Upload the .docx
    Or directly drag-and-drop the .docx into Google Drive.
    """
    if not req.answer.strip():
        raise HTTPException(400, "Answer cannot be empty")

    try:
        docx_bytes = answer_to_docx(req.answer, req.question, req.sources)
    except RuntimeError as e:
        raise HTTPException(503, str(e))
    except Exception as e:
        raise HTTPException(500, f"DOCX generation failed: {e}")

    filename = "DocuMindAI_Answer.docx"
    if req.question:
        safe = re.sub(r'[^\w\s-]', '', req.question.strip())[:40]
        filename = f"{safe.replace(' ', '_')}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
