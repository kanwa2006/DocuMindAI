"""
Phase 11 — Teacher Table Extraction Service.

Routing:
  native PDF   → Docling (fastest, preserves structure)
  scanned/image → PaddleOCR PPStructure
  fallback      → Docling → if 0 tables → PaddleOCR → if 0 tables → empty list
"""
import io
import re
import csv
import logging
from typing import Any, Dict, List, Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class TableExtractor:
    def __init__(self) -> None:
        self._converter = None  # lazy-loaded

    def _get_converter(self):
        if self._converter is None:
            from docling.document_converter import DocumentConverter
            self._converter = DocumentConverter()
        return self._converter

    async def extract_tables(
        self,
        file_path: str,
        is_scanned: bool = False,
    ) -> List[Dict[str, Any]]:
        """
        Returns list of table dicts:
        {
          "table_index": int,
          "page": int | None,
          "rows": List[List[str]],
          "has_header": bool,
          "caption": str | None,
          "merged_cells": List[dict]
        }
        """
        if not Path(file_path).exists():
            logger.error("[table_extractor] File not found: %s", file_path)
            return []

        if is_scanned:
            tables = await self._extract_with_paddle(file_path)
        else:
            tables = await self._extract_with_docling(file_path)
            if not tables:
                logger.info("[table_extractor] Docling found 0 tables, trying PaddleOCR")
                tables = await self._extract_with_paddle(file_path)

        return tables

    async def _extract_with_docling(self, file_path: str) -> List[Dict]:
        import asyncio
        from functools import partial
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, partial(self._docling_sync, file_path))

    def _docling_sync(self, file_path: str) -> List[Dict]:
        try:
            converter = self._get_converter()
            result = converter.convert(file_path)
            tables = []
            for i, table in enumerate(result.document.tables):
                rows: List[List[str]] = []
                for row in table.data:
                    rows.append([cell.text.strip() if cell.text else "" for cell in row])

                page: Optional[int] = None
                try:
                    page = table.prov[0].page_no if table.prov else None
                except Exception:
                    pass

                caption: Optional[str] = None
                try:
                    caption = table.caption_text if hasattr(table, "caption_text") else None
                except Exception:
                    pass

                if rows:
                    tables.append({
                        "table_index": i,
                        "page": page,
                        "rows": rows,
                        "has_header": len(rows) > 1,
                        "caption": caption,
                        "merged_cells": [],
                    })
            logger.info("[table_extractor] Docling extracted %d tables from %s", len(tables), file_path)
            return tables
        except Exception as exc:
            logger.warning("[table_extractor] Docling extraction failed: %s", exc)
            return []

    async def _extract_with_paddle(self, file_path: str) -> List[Dict]:
        import asyncio
        from functools import partial
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, partial(self._paddle_sync, file_path))

    def _paddle_sync(self, file_path: str) -> List[Dict]:
        try:
            from paddleocr import PPStructure
            import cv2
            engine = PPStructure(table=True, ocr=True, show_log=False, lang="en")
            img = cv2.imread(file_path)
            if img is None:
                # Try converting PDF page to image first
                import fitz
                doc = fitz.open(file_path)
                if len(doc) == 0:
                    doc.close()
                    return []
                page = doc[0]
                pix = page.get_pixmap(dpi=150)
                import numpy as np
                img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
                if pix.n == 4:
                    img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)
                doc.close()

            result = engine(img)
            tables = []
            idx = 0
            for region in result:
                if region.get("type") == "table":
                    html = region.get("res", {}).get("html", "")
                    rows = self._parse_html_table(html)
                    if rows:
                        tables.append({
                            "table_index": idx,
                            "page": 1,
                            "rows": rows,
                            "has_header": True,
                            "caption": None,
                            "merged_cells": [],
                        })
                        idx += 1
            logger.info("[table_extractor] PaddleOCR extracted %d tables from %s", len(tables), file_path)
            return tables
        except Exception as exc:
            logger.warning("[table_extractor] PaddleOCR extraction failed: %s", exc)
            return []

    def _parse_html_table(self, html: str) -> List[List[str]]:
        rows: List[List[str]] = []
        for row_match in re.finditer(r"<tr[^>]*>(.*?)</tr>", html, re.DOTALL | re.IGNORECASE):
            cells = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", row_match.group(1), re.DOTALL | re.IGNORECASE)
            cleaned = [re.sub(r"<[^>]+>", "", c).strip() for c in cells]
            if any(cleaned):
                rows.append(cleaned)
        return rows

    # ── Export methods ────────────────────────────────────────────────────────

    def export_to_docx(self, table: Dict, output_path: str) -> str:
        from docx import Document as DocxDocument
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = DocxDocument()
        if table.get("caption"):
            p = doc.add_paragraph(table["caption"])
            p.runs[0].italic = True

        rows = table.get("rows", [])
        if not rows:
            doc.save(output_path)
            return output_path

        col_count = max(len(r) for r in rows)
        t = doc.add_table(rows=len(rows), cols=col_count)
        t.style = "Table Grid"

        for row_idx, row in enumerate(rows):
            for col_idx in range(col_count):
                cell_text = row[col_idx] if col_idx < len(row) else ""
                cell = t.cell(row_idx, col_idx)
                cell.text = cell_text
                if row_idx == 0 and table.get("has_header"):
                    para = cell.paragraphs[0]
                    if para.runs:
                        para.runs[0].bold = True
                    tc = cell._tc
                    tc_pr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "DBEAFE")
                    shd.set(qn("w:val"), "clear")
                    tc_pr.append(shd)

        doc.save(output_path)
        return output_path

    def export_to_csv(self, table: Dict) -> str:
        output = io.StringIO()
        writer = csv.writer(output, quoting=csv.QUOTE_ALL)
        for row in table.get("rows", []):
            writer.writerow(row)
        return output.getvalue()

    def export_to_html(self, table: Dict) -> str:
        rows = table.get("rows", [])
        has_header = table.get("has_header", False)
        lines = ['<table border="1" style="border-collapse:collapse;font-family:Arial,sans-serif;">']
        for row_idx, row in enumerate(rows):
            lines.append("  <tr>")
            tag = "th" if row_idx == 0 and has_header else "td"
            for cell in row:
                style = "padding:8px 12px;"
                if tag == "th":
                    style += "background:#DBEAFE;font-weight:600;"
                # Escape HTML special chars
                safe = (
                    str(cell)
                    .replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace('"', "&quot;")
                )
                lines.append(f'    <{tag} style="{style}">{safe}</{tag}>')
            lines.append("  </tr>")
        lines.append("</table>")
        return "\n".join(lines)


# ── Singleton ─────────────────────────────────────────────────────────────────

_instance: Optional[TableExtractor] = None


def get_table_extractor() -> TableExtractor:
    global _instance
    if _instance is None:
        _instance = TableExtractor()
    return _instance
