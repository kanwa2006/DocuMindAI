import os
import logging
from fpdf import FPDF
from docx import Document
from typing import Dict, Any

logger = logging.getLogger(__name__)

LOCAL_STORAGE_DIR = "/tmp/documind_storage/exports"
os.makedirs(LOCAL_STORAGE_DIR, exist_ok=True)

class PDFReport(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 15)
        self.cell(0, 10, 'DocuMindAI Enterprise Report', border=False, align='C')
        self.ln(20)
        
    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', border=False, align='C')

class ExportService:
    @staticmethod
    def generate_pdf(payload: Dict[str, Any], filename: str) -> str:
        pdf = PDFReport()
        pdf.add_page()
        
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, payload.get("title", "Generated Report"), ln=True)
        pdf.ln(5)
        
        if "confidence_score" in payload:
            pdf.set_font("helvetica", "I", 10)
            pdf.cell(0, 10, f"Confidence Score: {payload['confidence_score']:.1%}", ln=True)
            pdf.ln(5)
            
        pdf.set_font("helvetica", "", 12)
        pdf.multi_cell(0, 8, payload.get("content", ""))
        pdf.ln(10)
        
        if payload.get("evidence"):
            pdf.set_font("helvetica", "B", 14)
            pdf.cell(0, 10, "Citation Evidence", ln=True)
            for i, chunk in enumerate(payload["evidence"]):
                pdf.set_font("helvetica", "B", 10)
                pdf.cell(0, 8, f"[{i+1}] {chunk.get('filename')} (Page {chunk.get('page_number')})", ln=True)
                pdf.set_font("helvetica", "I", 10)
                pdf.multi_cell(0, 6, f'"{chunk.get("text_content")}"')
                pdf.ln(4)
                
        if payload.get("diagnostics"):
            pdf.add_page()
            pdf.set_font("helvetica", "B", 14)
            pdf.cell(0, 10, "Retrieval Diagnostics", ln=True)
            pdf.set_font("courier", "", 10)
            for k, v in payload["diagnostics"].items():
                pdf.cell(0, 6, f"{k}: {v}", ln=True)
                
        file_path = os.path.join(LOCAL_STORAGE_DIR, filename)
        pdf.output(file_path)
        return file_path

    @staticmethod
    def generate_docx(payload: Dict[str, Any], filename: str) -> str:
        doc = Document()
        doc.add_heading('DocuMindAI Enterprise Report', 0)
        
        doc.add_heading(payload.get("title", "Generated Report"), level=1)
        
        if "confidence_score" in payload:
            p = doc.add_paragraph()
            p.add_run(f"Confidence Score: {payload['confidence_score']:.1%}").italic = True
            
        doc.add_paragraph(payload.get("content", ""))
        
        if payload.get("evidence"):
            doc.add_heading('Citation Evidence', level=2)
            for i, chunk in enumerate(payload["evidence"]):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"[{i+1}] {chunk.get('filename')} (Page {chunk.get('page_number')}): ").bold = True
                p.add_run(f'"{chunk.get("text_content")}"').italic = True
                
        if payload.get("diagnostics"):
            doc.add_page_break()
            doc.add_heading('Retrieval Diagnostics', level=2)
            for k, v in payload["diagnostics"].items():
                doc.add_paragraph(f"{k}: {v}")
                
        file_path = os.path.join(LOCAL_STORAGE_DIR, filename)
        doc.save(file_path)
        return file_path

    @staticmethod
    def generate_exam_docx(payload: Dict[str, Any], filename: str) -> str:
        """
        PHASE 5: True Export Engine
        Handles the strict Exam JSON hierarchy and renders printable DOCX formatting.
        """
        doc = Document()
        
        # Exam Header
        doc.add_heading(payload.get("title", "Untitled Exam"), 0)
        if payload.get("description"):
            doc.add_paragraph(payload.get("description")).italic = True
            
        content = payload.get("content", {}).get("sections", [])
        total_marks = sum(
            sum(q.get("marks", 0) for q in sec.get("questions", []))
            for sec in content
        )
        
        p = doc.add_paragraph()
        p.add_run(f"Total Marks: {total_marks}").bold = True
        
        # Sections
        for s_idx, section in enumerate(content):
            doc.add_heading(f"Section {chr(65+s_idx)}: {section.get('title', '')}", level=1)
            
            for q_idx, q in enumerate(section.get("questions", [])):
                p_q = doc.add_paragraph()
                p_q.add_run(f"Q{q_idx + 1}. ").bold = True
                p_q.add_run(q.get("text", ""))
                p_q.add_run(f" [{q.get('marks', 0)} marks]").bold = True
                
                sub_qs = q.get("sub_questions", [])
                for sq in sub_qs:
                    p_sq = doc.add_paragraph(style='List Bullet 2')
                    p_sq.add_run(f"{sq.get('label', '')}. ").bold = True
                    p_sq.add_run(sq.get("text", ""))
                    p_sq.add_run(f" [{sq.get('marks', 0)} marks]").italic = True
                    
        # Optional: Answer Key separated by page break
        doc.add_page_break()
        doc.add_heading("Answer Key & Marking Scheme", level=1)
        for s_idx, section in enumerate(content):
            for q_idx, q in enumerate(section.get("questions", [])):
                if q.get("answer_key") or q.get("rubric"):
                    doc.add_heading(f"Q{q_idx + 1} Key", level=3)
                    if q.get("answer_key"):
                        doc.add_paragraph("Answer: " + q.get("answer_key"))
                    if q.get("rubric"):
                        p_r = doc.add_paragraph("Rubric: ")
                        p_r.add_run(q.get("rubric")).italic = True
                        
                sub_qs = q.get("sub_questions", [])
                for sq in sub_qs:
                    if sq.get("answer_key") or sq.get("rubric"):
                        p_sk = doc.add_paragraph(style='List Bullet 2')
                        p_sk.add_run(f"{sq.get('label', '')}. ")
                        p_sk.add_run(sq.get("answer_key", ""))

        file_path = os.path.join(LOCAL_STORAGE_DIR, filename)
        doc.save(file_path)
        return file_path
