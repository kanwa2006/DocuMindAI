import uuid
import json
from io import BytesIO
from typing import List, Any, Optional
from fastapi import APIRouter, Depends, HTTPException, Body
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select
from sqlalchemy.orm import selectinload
from pydantic import BaseModel
import logging

from app.db.session import get_db, AsyncSessionLocal
from app.core.auth import get_current_user
from app.core.workspace import resolve_workspace_id
from app.models.export_job import ExportJob
from app.models.legal import Contract, Clause
from app.models.chat import ChatSession
from app.schemas.export import ExportJobCreate, ExportJobResponse
from app.workers.tasks.export_tasks import process_export_job
from app.services.export_engine import export_engine
from app.services.audit_export import generate_session_audit_report


class SessionExportOptions(BaseModel):
    include_citations: bool = True
    include_confidence: bool = False
    include_disclaimers: bool = True


_WORKSPACE_DISCLAIMERS = {
    "legal": (
        "DISCLAIMER: This analysis is AI-generated for informational purposes only. "
        "It does not constitute legal advice. Always consult a qualified legal professional."
    ),
    "finance": (
        "DISCLAIMER: All figures are AI-extracted. Verify all numbers against original "
        "source documents before any financial, tax, or legal use."
    ),
}

router = APIRouter()
logger = logging.getLogger(__name__)

@router.get("/legal/{contract_id}/docx")
async def export_legal_contract(contract_id: str, current_user: dict = Depends(get_current_user)):
    """
    PHASE 4: REAL EXPORT ENGINE
    Generates a DOCX format redline review for Legal Workspaces.
    """
    async with AsyncSessionLocal() as db:
        # Fetch contract with clauses and redlines
        stmt = select(Contract).where(Contract.id == contract_id).options(
            selectinload(Contract.clauses).selectinload(Clause.redlines)
        )
        contract = (await db.execute(stmt)).scalar_one_or_none()
        
        if not contract:
            raise HTTPException(status_code=404, detail="Contract not found")
            
        # Format data for engine
        clauses_data = []
        for clause in contract.clauses:
            c_data = {
                "clause": {
                    "section_name": clause.section_name,
                    "original_text": clause.original_text,
                    "risk_level": clause.risk_level,
                    "compliance_notes": clause.compliance_notes
                },
                "redlines": [{"suggested_text": r.suggested_text} for r in clause.redlines]
            }
            clauses_data.append(c_data)

    # Generate DOCX in memory
    try:
        file_stream = export_engine.generate_legal_redline_docx(contract.title, clauses_data)
        
        # PHASE 9: SECURITY HARDENING
        # Ensure headers prevent caching of sensitive files
        headers = {
            'Content-Disposition': f'attachment; filename="redline_{contract.id}.docx"',
            'Cache-Control': 'no-store, no-cache, must-revalidate, max-age=0',
        }
        return StreamingResponse(file_stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
        
    except Exception as e:
        logger.error(f"DOCX Export Failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate DOCX file.")

@router.post("", response_model=ExportJobResponse)
async def create_export_job(
    request: ExportJobCreate,
    current_user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """
    Trigger an asynchronous background export job.
    Enforces strict tenant isolation.
    """
    job = ExportJob(
        owner_id=uuid.UUID(current_user["id"]),
        workspace_id=resolve_workspace_id(current_user["workspace_id"]),
        format=request.format,
        export_type=request.export_type,
        payload=request.payload
    )
    db.add(job)
    await db.commit()
    await db.refresh(job)
    
    # Trigger Celery Task
    contract_id = request.payload.get("contract_id", "") if request.payload else ""
    process_export_job.delay(contract_id, str(job.id))
    
    return job

@router.get("", response_model=List[ExportJobResponse])
async def list_exports(
    current_user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
) -> Any:
    """Fetch all export jobs for the current tenant's workspace."""
    stmt = select(ExportJob).where(
        ExportJob.workspace_id == resolve_workspace_id(current_user["workspace_id"])
    ).order_by(ExportJob.created_at.desc())
    
    result = await db.execute(stmt)
    return result.scalars().all()

@router.get("/{job_id}", response_model=ExportJobResponse)
async def get_export_job(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
) -> Any:
    """Poll export job status."""
    try:
        job_uuid = uuid.UUID(job_id)
    except ValueError:
        raise HTTPException(status_code=422, detail="Invalid Job ID format.")

    stmt = select(ExportJob).where(
        ExportJob.id == job_uuid,
        ExportJob.workspace_id == resolve_workspace_id(current_user["workspace_id"])
    )
    result = await db.execute(stmt)
    job = result.scalar_one_or_none()

    if not job:
        raise HTTPException(status_code=404, detail="Export job not found or access denied.")

    return job


# ─── Session Export Endpoints (Phase 8) ──────────────────────────────────────

@router.post("/sessions/{session_id}/pdf")
async def export_session_pdf(
    session_id: str,
    opts: SessionExportOptions = Body(default_factory=SessionExportOptions),
    current_user: dict = Depends(get_current_user),
):
    """Export a chat session as PDF. Fetches messages and generates formatted PDF."""
    async with AsyncSessionLocal() as db:
        stmt = (
            select(ChatSession)
            .where(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.owner_id == uuid.UUID(current_user["id"]),
            )
            .options(selectinload(ChatSession.messages))
        )
        session = (await db.execute(stmt)).scalar_one_or_none()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found.")

    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title (H1)
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, session.title[:80], new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(2)

        # Workspace badge
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, f"{session.workspace_type.upper()} WORKSPACE", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

        # Disclaimer
        if opts.include_disclaimers and session.workspace_type in _WORKSPACE_DISCLAIMERS:
            pdf.set_font("Helvetica", "I", 9)
            pdf.set_text_color(120, 80, 0)
            pdf.multi_cell(0, 5, _WORKSPACE_DISCLAIMERS[session.workspace_type])
            pdf.set_text_color(0, 0, 0)
            pdf.ln(4)

        # Messages
        for msg in session.messages:
            if msg.role == "user":
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(59, 130, 246)
                pdf.cell(0, 6, "You:", new_x="LMARGIN", new_y="NEXT")
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Helvetica", "", 11)
                pdf.set_fill_color(245, 245, 245)
                pdf.multi_cell(0, 6, msg.content, fill=True)
                pdf.ln(3)

            elif msg.role == "assistant":
                answer_text = msg.content
                evidence: list = []
                confidence: float = 0.0
                try:
                    parsed = json.loads(msg.content)
                    answer_text = parsed.get("answer", msg.content)
                    evidence = parsed.get("evidence", [])
                    confidence = float(parsed.get("confidence_score", 0.0))
                except Exception:
                    pass

                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(22, 163, 74)
                pdf.cell(0, 6, "DocuMindAI:", new_x="LMARGIN", new_y="NEXT")
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Helvetica", "", 11)
                clean = (
                    answer_text.replace("**", "").replace("*", "")
                    .replace("###", "").replace("##", "").replace("#", "")
                    .replace("`", "")
                )
                pdf.multi_cell(0, 6, clean)

                if opts.include_confidence and confidence > 0:
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.set_text_color(100, 100, 100)
                    pdf.cell(0, 5, f"Confidence: {int(confidence * 100)}%", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(0, 0, 0)

                if opts.include_citations and evidence:
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.set_text_color(100, 100, 100)
                    for i, chunk in enumerate(evidence[:6], 1):
                        fn = chunk.get("filename", "Unknown")
                        pg = chunk.get("page_number", "?")
                        pdf.cell(0, 5, f"[{i}] {fn}, p.{pg}", new_x="LMARGIN", new_y="NEXT")
                    pdf.set_text_color(0, 0, 0)

                pdf.ln(3)

        # Footer on current page
        pdf.set_y(-15)
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(150, 150, 150)
        from datetime import datetime as _dt
        pdf.cell(0, 10, f"Generated by DocuMindAI | grounded answers | {_dt.utcnow().strftime('%Y-%m-%d')}", align="C")

        pdf_bytes = pdf.output()
        safe = "".join(c if c.isalnum() or c in "- _" else "_" for c in session.title)[:50]
        return StreamingResponse(
            BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{safe}.pdf"',
                "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            },
        )
    except Exception as e:
        logger.error(f"Session PDF export failed: {e}")
        raise HTTPException(status_code=500, detail="PDF generation failed.")


@router.get("/sessions/{session_id}/audit-report")
async def export_session_audit_report(
    session_id: str,
    format: str = "pdf",
    current_user: dict = Depends(get_current_user),
):
    """
    Phase 18-C — Audit Trail Export.
    Generates a structured audit report (PDF or DOCX) for the given session.

    Plan gating:
      trial        → 402 with upgrade_required: true
      professional → PDF only
      enterprise   → PDF + DOCX
    """
    plan = (current_user.get("plan") or "trial").lower()

    if plan == "trial":
        raise HTTPException(
            status_code=402,
            detail={"error": "upgrade_required", "upgrade_required": True,
                    "message": "Audit Trail Export is available on Professional and Enterprise plans."},
        )

    if plan == "professional" and format == "docx":
        raise HTTPException(
            status_code=402,
            detail={"error": "upgrade_required", "upgrade_required": True,
                    "message": "DOCX export is available on the Enterprise plan."},
        )

    if format not in ("pdf", "docx"):
        raise HTTPException(status_code=422, detail="format must be 'pdf' or 'docx'.")

    async with AsyncSessionLocal() as db:
        stmt = (
            select(ChatSession)
            .where(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.owner_id == uuid.UUID(current_user["id"]),
            )
            .options(selectinload(ChatSession.messages))
        )
        session = (await db.execute(stmt)).scalar_one_or_none()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found.")

    try:
        buf = generate_session_audit_report(session, current_user, fmt=format)  # type: ignore[arg-type]

        safe_title = "".join(c if c.isalnum() or c in "- _" else "_" for c in session.title)[:50]

        # Track analytics event
        try:
            logger.info(
                "export_downloaded event: session=%s format=audit_%s user=%s",
                session_id, format, current_user.get("id"),
            )
        except Exception:
            pass

        if format == "pdf":
            return StreamingResponse(
                buf,
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f'attachment; filename="audit_{safe_title}.pdf"',
                    "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
                },
            )
        else:
            return StreamingResponse(
                buf,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f'attachment; filename="audit_{safe_title}.docx"',
                    "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
                },
            )
    except Exception as e:
        logger.error(f"Audit report generation failed: {e}")
        raise HTTPException(status_code=500, detail="Audit report generation failed.")


@router.post("/sessions/{session_id}/docx")
async def export_session_docx(
    session_id: str,
    opts: SessionExportOptions = Body(default_factory=SessionExportOptions),
    current_user: dict = Depends(get_current_user),
):
    """Export a chat session as DOCX. Fetches messages and generates formatted Word document."""
    async with AsyncSessionLocal() as db:
        stmt = (
            select(ChatSession)
            .where(
                ChatSession.id == uuid.UUID(session_id),
                ChatSession.owner_id == uuid.UUID(current_user["id"]),
            )
            .options(selectinload(ChatSession.messages))
        )
        session = (await db.execute(stmt)).scalar_one_or_none()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found.")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Heading 1: session title
        h = doc.add_heading(session.title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Workspace badge in header section
        section = doc.sections[0]
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(f"{session.workspace_type.upper()} WORKSPACE")
        hr.font.size = Pt(9)
        hr.font.color.rgb = RGBColor(100, 100, 100)

        # Disclaimer
        if opts.include_disclaimers and session.workspace_type in _WORKSPACE_DISCLAIMERS:
            dp = doc.add_paragraph(_WORKSPACE_DISCLAIMERS[session.workspace_type])
            dp.runs[0].font.size = Pt(9)
            dp.runs[0].italic = True
            dp.runs[0].font.color.rgb = RGBColor(120, 80, 0)

        doc.add_paragraph()

        # Messages
        for msg in session.messages:
            if msg.role == "user":
                p = doc.add_paragraph()
                label = p.add_run("You: ")
                label.bold = True
                label.font.color.rgb = RGBColor(59, 130, 246)
                p.add_run(msg.content)

            elif msg.role == "assistant":
                answer_text = msg.content
                evidence: list = []
                confidence: float = 0.0
                try:
                    parsed = json.loads(msg.content)
                    answer_text = parsed.get("answer", msg.content)
                    evidence = parsed.get("evidence", [])
                    confidence = float(parsed.get("confidence_score", 0.0))
                except Exception:
                    pass

                p = doc.add_paragraph()
                label = p.add_run("DocuMindAI: ")
                label.bold = True
                label.font.color.rgb = RGBColor(22, 163, 74)
                p.add_run(answer_text)

                if opts.include_confidence and confidence > 0:
                    cp = doc.add_paragraph()
                    cr = cp.add_run(f"Confidence: {int(confidence * 100)}%")
                    cr.font.size = Pt(9)
                    cr.italic = True

                if opts.include_citations and evidence:
                    for i, chunk in enumerate(evidence[:6], 1):
                        fn = chunk.get("filename", "Unknown")
                        pg = chunk.get("page_number", "?")
                        cp = doc.add_paragraph()
                        cr = cp.add_run(f"[{i}] {fn}, p.{pg}")
                        cr.font.size = Pt(9)
                        cr.italic = True

            # Horizontal separator between message pairs
            doc.add_paragraph("─" * 50)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        safe = "".join(c if c.isalnum() or c in "- _" else "_" for c in session.title)[:50]
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{safe}.docx"',
                "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
            },
        )
    except Exception as e:
        logger.error(f"Session DOCX export failed: {e}")
        raise HTTPException(status_code=500, detail="DOCX generation failed.")
